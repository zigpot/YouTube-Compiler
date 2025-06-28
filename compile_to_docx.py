import os
import sys
from docx import Document

if len(sys.argv) < 2:
    print("Usage: python script.py <input_directory> [output_directory]")
    sys.exit(1)

input_dir = sys.argv[1]
output_dir = sys.argv[2] if len(sys.argv) > 2 else input_dir

if not os.path.isdir(input_dir):
    print(f"Error: {input_dir} is not a valid directory.")
    sys.exit(1)

doc = Document()
doc.add_heading("Compiled Text", 0)

for filename in sorted(os.listdir(input_dir)):
    if filename.endswith(".txt"):
        file_path = os.path.join(input_dir, filename)
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            doc.add_heading(filename[:-4], level=2)
            doc.add_paragraph(text)
            doc.add_paragraph("")

output_path = os.path.join(output_dir, "compiled.docx")
doc.save(output_path)
print(f"Saved to {output_path}")
